"""Helpers for rendering .docx templates and converting to PDF."""

from __future__ import annotations

import io
import os
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from typing import Any, Iterable

import httpx
from docx import Document
from docx.enum.text import WD_BREAK
from docxtpl import DocxTemplate
from jinja2 import Environment, StrictUndefined
from jinja2.exceptions import UndefinedError, TemplateSyntaxError, TemplateError


GOTENBERG_URL = os.environ.get("GOTENBERG_URL", "http://gotenberg:3000")
MAX_TEMPLATE_BYTES = 5 * 1024 * 1024
MAX_ITEMS_PER_EXPORT = 500


class TemplateRenderError(Exception):
  """Raised when a docx template cannot be rendered (missing var, bad syntax, ...)."""


class PdfConversionError(Exception):
  """Raised when Gotenberg cannot convert the docx to PDF."""


# OOXML namespaces (Word comments + body).
_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
_NS_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

ET.register_namespace("w", _NS_W)


def _qn(tag: str) -> str:
  return f"{{{_NS_W}}}{tag}"


def _localname(tag: str) -> str:
  if "}" in tag:
    return tag.rsplit("}", 1)[-1]
  return tag


def _attr_id(elem: ET.Element) -> str | None:
  """Resolve `w:id` / namespaced id attribute."""
  for key, val in elem.attrib.items():
    if key.endswith("}id") or key == "id":
      return val
  return None


def _parent_map(root: ET.Element) -> dict[ET.Element, ET.Element]:
  m: dict[ET.Element, ET.Element] = {}
  for parent in root.iter():
    for child in parent:
      m[child] = parent
  return m


def _dfs_preorder(elem: ET.Element) -> list[ET.Element]:
  out: list[ET.Element] = [elem]
  for child in elem:
    out.extend(_dfs_preorder(child))
  return out


def _tree_depth(elem: ET.Element, parents: dict[ET.Element, ET.Element]) -> int:
  d = 0
  cur: ET.Element | None = elem
  while cur is not None:
    d += 1
    cur = parents.get(cur)
  return d


def _parse_comments_xml(data: bytes) -> dict[str, dict[str, Any]]:
  """Map comment id -> {text, para_id}."""
  root = ET.fromstring(data)
  out: dict[str, dict[str, Any]] = {}
  for cm in root:
    if _localname(cm.tag) != "comment":
      continue
    cid = _attr_id(cm)
    if cid is None:
      continue
    para_id: str | None = None
    text_parts: list[str] = []
    for node in cm.iter():
      if _localname(node.tag) == "t" and node.text:
        text_parts.append(node.text)
      if _localname(node.tag) == "p":
        pid = node.get(f"{{{_NS_W14}}}paraId")
        if pid:
          para_id = pid
    out[cid] = {"text": "".join(text_parts).strip(), "para_id": para_id}
  return out


def _parse_comments_extended_resolved(data: bytes) -> set[str]:
  """Para ids marked resolved (done) in commentsExtended.xml."""
  root = ET.fromstring(data)
  resolved: set[str] = set()
  for el in root.iter():
    if _localname(el.tag) != "commentEx":
      continue
    para_id = el.get(f"{{{_NS_W15}}}paraId")
    done = el.get(f"{{{_NS_W15}}}done")
    if para_id and done == "1":
      resolved.add(para_id)
  return resolved


def _pair_comment_markers(body: ET.Element, parents: dict[ET.Element, ET.Element]) -> list[tuple[str, ET.Element, ET.Element]]:
  """Pair commentRangeStart with commentRangeEnd by id using document-order stack."""
  pairs: list[tuple[str, ET.Element, ET.Element]] = []
  stack: list[tuple[str, ET.Element]] = []
  for el in _dfs_preorder(body):
    ln = _localname(el.tag)
    if ln == "commentRangeStart":
      cid = _attr_id(el)
      if cid is not None:
        stack.append((cid, el))
    elif ln == "commentRangeEnd":
      cid = _attr_id(el)
      if cid is None:
        continue
      if not stack or stack[-1][0] != cid:
        continue
      _, start_el = stack.pop()
      pairs.append((cid, start_el, el))
  return pairs


def _collect_range_inclusive(body: ET.Element, start_el: ET.Element, end_el: ET.Element) -> list[ET.Element]:
  """All elements from start_el through end_el in preorder DFS under body."""
  acc: list[ET.Element] = []
  inside = False
  for el in _dfs_preorder(body):
    if el is start_el:
      inside = True
    if inside:
      acc.append(el)
      if el is end_el:
        break
  return acc


def _make_run_with_text(text: str) -> ET.Element:
  r_el = ET.Element(_qn("r"))
  t_el = ET.Element(_qn("t"))
  if text.startswith(" ") or text.endswith(" ") or "\n" in text:
    t_el.set(_XML_SPACE, "preserve")
  t_el.text = text
  r_el.append(t_el)
  return r_el


def _remove_elements_deepest_first(elems: Iterable[ET.Element], parents: dict[ET.Element, ET.Element]) -> None:
  sorted_elems = sorted(elems, key=lambda e: _tree_depth(e, parents), reverse=True)
  for el in sorted_elems:
    par = parents.get(el)
    if par is not None:
      try:
        par.remove(el)
      except ValueError:
        pass


def _remove_comment_reference_runs(root: ET.Element, cid: str, parents: dict[ET.Element, ET.Element]) -> None:
  """Remove w:r runs that only contain w:commentReference for this id."""
  to_remove: list[ET.Element] = []
  for el in root.iter():
    if _localname(el.tag) != "commentReference":
      continue
    if _attr_id(el) != cid:
      continue
    par = parents.get(el)
    if par is None or _localname(par.tag) != "r":
      continue
    only_ref = True
    for child in par:
      if _localname(child.tag) != "commentReference":
        only_ref = False
        break
    if only_ref:
      to_remove.append(par)
  _remove_elements_deepest_first(to_remove, parents)


def _strip_all_comment_markup(root: ET.Element) -> None:
  """Best-effort cleanup for any remaining comment tags in document.xml."""
  parents = _parent_map(root)
  # Remove all range markers globally.
  markers: list[ET.Element] = []
  for el in root.iter():
    ln = _localname(el.tag)
    if ln in ("commentRangeStart", "commentRangeEnd"):
      markers.append(el)
  _remove_elements_deepest_first(markers, parents)

  # Remove all comment references. Prefer removing the whole run when it only
  # contains commentReference nodes; otherwise remove just the node.
  parents = _parent_map(root)
  refs = [el for el in root.iter() if _localname(el.tag) == "commentReference"]
  to_remove_runs: list[ET.Element] = []
  to_remove_nodes: list[ET.Element] = []
  for ref in refs:
    par = parents.get(ref)
    if par is None:
      continue
    if _localname(par.tag) == "r":
      only_refs = all(_localname(ch.tag) == "commentReference" for ch in par)
      if only_refs:
        to_remove_runs.append(par)
      else:
        to_remove_nodes.append(ref)
    else:
      to_remove_nodes.append(ref)
  _remove_elements_deepest_first(to_remove_runs, parents)
  parents = _parent_map(root)
  _remove_elements_deepest_first(to_remove_nodes, parents)


_COMMENT_PART_NAMES_LOWER = frozenset({
  "word/comments.xml",
  "word/commentsextended.xml",
  "word/commentsids.xml",
  "word/commentsextensible.xml",
})


def _patch_content_types_and_rels(docx_bytes: bytes) -> bytes:
  """Remove Overrides for stripped comment parts and Relationships pointing at them."""
  buf = io.BytesIO(docx_bytes)
  out_buf = io.BytesIO()

  with zipfile.ZipFile(buf, "r") as zin:
    names_lower = {n.lower(): n for n in zin.namelist()}
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
      for info in zin.infolist():
        name = info.filename
        ln = name.lower()

        if ln == "[content_types].xml":
          data = zin.read(name)
          try:
            root = ET.fromstring(data)
            children = list(root)
            for child in children:
              tag = _localname(child.tag)
              part = (child.get("PartName") or "").lstrip("/").lower()
              ctype = (child.get("ContentType") or "").lower()
              if tag == "Override":
                # Strip any comments part override (comments.xml, commentsExtended.xml,
                # commentsExtensible.xml, etc).
                if "comment" in part or "comment" in ctype:
                  root.remove(child)
              elif tag == "Default":
                # Some files declare comment content-types on Default entries.
                if "comment" in ctype:
                  root.remove(child)
            data = ET.tostring(root, encoding="utf-8", xml_declaration=True, default_namespace=None)
          except ET.ParseError:
            pass
          zout.writestr(info, data)
          continue

        if ln == "word/_rels/document.xml.rels":
          data = zin.read(name)
          try:
            root = ET.fromstring(data)
            children = list(root)
            for child in children:
              if child.tag.endswith("Relationship"):
                tgt = (child.get("Target") or "").lower()
                rel_type = (child.get("Type") or "").lower()
                # Remove all comment-related relationships, including
                # commentsExtensible/commentsExtended variants.
                if "comment" in rel_type or "comment" in tgt:
                  root.remove(child)
            data = ET.tostring(root, encoding="utf-8", xml_declaration=True, default_namespace=None)
          except ET.ParseError:
            pass
          zout.writestr(info, data)
          continue

        zout.writestr(info, zin.read(name))

  return out_buf.getvalue()


def _apply_comment_placeholders(template_bytes: bytes) -> bytes:
  """Replace highlighted regions with Jinja from Word comments; strip comment XML parts.

  If there are no comments in the package, returns template_bytes unchanged.
  """
  try:
    with zipfile.ZipFile(io.BytesIO(template_bytes), "r") as zf:
      names = set(zf.namelist())
      by_lower = {n.replace("\\", "/").lower(): n for n in names}
      if "word/comments.xml" not in by_lower:
        return template_bytes

      comments_raw = zf.read(by_lower["word/comments.xml"])
      comments_map = _parse_comments_xml(comments_raw)

      resolved_para: set[str] = set()
      if "word/commentsextended.xml" in by_lower:
        resolved_para = _parse_comments_extended_resolved(zf.read(by_lower["word/commentsextended.xml"]))

      doc_xml = zf.read("word/document.xml")
      root = ET.fromstring(doc_xml)
      body = root.find(_qn("body"))
      if body is None:
        return template_bytes

      parents = _parent_map(root)
      pairs = _pair_comment_markers(body, parents)

      for cid, start_el, end_el in pairs:
        meta = comments_map.get(cid, {})
        text = (meta.get("text") or "").strip()
        para_id = meta.get("para_id")
        is_resolved = bool(para_id and para_id in resolved_para)

        if is_resolved or not text:
          # Keep highlighted text; drop only range markers and comment references.
          try:
            ps = parents.get(start_el)
            if ps is not None:
              ps.remove(start_el)
          except ValueError:
            pass
          parents = _parent_map(root)
          try:
            pe = parents.get(end_el)
            if pe is not None:
              pe.remove(end_el)
          except ValueError:
            pass
          parents = _parent_map(root)
          _remove_comment_reference_runs(root, cid, parents)
          parents = _parent_map(root)
          continue

        rng = _collect_range_inclusive(body, start_el, end_el)
        new_run = _make_run_with_text(text)
        par_s = parents.get(start_el)
        if par_s is None:
          continue
        try:
          idx = list(par_s).index(start_el)
        except ValueError:
          continue
        par_s.insert(idx, new_run)
        parents = _parent_map(root)
        _remove_elements_deepest_first(rng, parents)
        parents = _parent_map(root)
        _remove_comment_reference_runs(root, cid, parents)
        parents = _parent_map(root)

      # Defensive pass: ensure no comment markup survives in document.xml.
      _strip_all_comment_markup(root)

      doc_out = ET.tostring(root, encoding="utf-8", xml_declaration=True)

      buf = io.BytesIO()
      with zipfile.ZipFile(io.BytesIO(template_bytes), "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
          doc_written = False
          for info in zin.infolist():
            fn = info.filename.replace("\\", "/")
            fl = fn.lower()
            if fl in _COMMENT_PART_NAMES_LOWER:
              continue
            if fl == "word/document.xml":
              zout.writestr(info, doc_out)
              doc_written = True
            else:
              zout.writestr(info, zin.read(info.filename))
          if not doc_written:
            zout.writestr("word/document.xml", doc_out)

      merged = buf.getvalue()
      return _patch_content_types_and_rels(merged)

  except Exception as exc:
    raise TemplateRenderError(f"Failed to read template comments: {exc}") from exc


def validate_docx_bytes(content: bytes) -> None:
  """Reject anything that isn't a real .docx zip with a word/document.xml entry."""
  if not content:
    raise ValueError("Empty file")
  if len(content) > MAX_TEMPLATE_BYTES:
    raise ValueError(f"Template exceeds {MAX_TEMPLATE_BYTES // 1024 // 1024} MB limit")
  try:
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
      names = set(zf.namelist())
  except zipfile.BadZipFile as exc:
    raise ValueError("File is not a valid .docx (must be an Office Open XML zip)") from exc
  if "word/document.xml" not in names:
    raise ValueError("Missing word/document.xml — not a Word document")


def _format_dt(value: Any) -> str:
  if isinstance(value, datetime):
    if value.tzinfo is None:
      value = value.replace(tzinfo=timezone.utc)
    return value.strftime("%Y-%m-%d %H:%M:%S %Z").strip()
  return str(value or "")


def build_context(
  items: list[dict],
  section: dict,
  account: dict,
  exported_by: str,
) -> dict:
  """Assemble the variables made available to the template."""
  normalised_items: list[dict] = []
  for it in items:
    normalised_items.append({
      "id": it.get("id", ""),
      "name": it.get("name", ""),
      "data": it.get("data") or {},
      "created_at": _format_dt(it.get("created_at")),
    })

  context: dict[str, Any] = {
    "items": normalised_items,
    "section": {
      "slug": section.get("slug", ""),
      "label": section.get("label", ""),
      "detail": section.get("detail", ""),
    },
    "account": {
      "id": account.get("id", ""),
      "name": account.get("name", ""),
    },
    "exported_at": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC"),
    "exported_by": exported_by or "",
  }

  # Convenience aliases for single-item templates.
  first = normalised_items[0] if normalised_items else {"name": "", "data": {}, "created_at": ""}
  context.setdefault("item", first)
  context.setdefault("name", first["name"])
  context.setdefault("data", first["data"])
  context.setdefault("created_at", first["created_at"])
  return context


def render_template(template_bytes: bytes, context: dict) -> bytes:
  """Render the docx template against the context.

  Uses Jinja's StrictUndefined so typos surface as TemplateRenderError instead of
  silently rendering as empty strings.
  """
  template_bytes = _apply_comment_placeholders(template_bytes)
  doc = DocxTemplate(io.BytesIO(template_bytes))
  jinja_env = Environment(undefined=StrictUndefined, autoescape=False)
  try:
    doc.render(context, jinja_env=jinja_env)
  except UndefinedError as exc:
    raise TemplateRenderError(f"Unknown variable in template: {exc}") from exc
  except TemplateSyntaxError as exc:
    raise TemplateRenderError(f"Template syntax error: {exc.message} (line {exc.lineno})") from exc
  except TemplateError as exc:
    raise TemplateRenderError(f"Template error: {exc}") from exc
  except Exception as exc:
    # docxtpl raises plain ValueError when a token is split across runs, etc.
    raise TemplateRenderError(f"Failed to render template: {exc}") from exc

  out = io.BytesIO()
  doc.save(out)
  return out.getvalue()


def docx_to_pdf(docx_bytes: bytes, filename: str = "document.docx") -> bytes:
  """Convert docx bytes to PDF bytes using Gotenberg's LibreOffice route."""
  url = f"{GOTENBERG_URL.rstrip('/')}/forms/libreoffice/convert"
  files = {
    "files": (
      filename,
      docx_bytes,
      "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
  }
  try:
    with httpx.Client(timeout=60.0) as client:
      resp = client.post(url, files=files)
  except httpx.HTTPError as exc:
    raise PdfConversionError(f"PDF service unreachable: {exc}") from exc

  if resp.status_code != 200:
    detail = resp.text.strip() or f"HTTP {resp.status_code}"
    raise PdfConversionError(f"PDF service error: {detail}")
  return resp.content


def _safe_field_keys(schema: Any) -> list[tuple[str, str]]:
  """Return (key, label) pairs for a section's schema, deduped, in declared order."""
  if not isinstance(schema, dict):
    return []
  fields = schema.get("fields") or []
  seen: set[str] = set()
  out: list[tuple[str, str]] = []
  for field in fields:
    if not isinstance(field, dict):
      continue
    key = str(field.get("key") or "").strip()
    if not key or key in seen:
      continue
    seen.add(key)
    label = str(field.get("label") or key)
    out.append((key, label))
  return out


def generate_starter_template(account_name: str, sections: Iterable[dict]) -> bytes:
  """Build a starter .docx that lists every placeholder available for this account."""
  doc = Document()
  doc.add_heading(f"Export template — {account_name}", level=0)

  intro = doc.add_paragraph()
  intro.add_run(
    "This document lists every placeholder you can use in your export template. "
    "Copy and paste into a new Word document, keep what you need, delete what you don't."
  )

  doc.add_heading("Quick reference", level=1)
  syntax_lines = [
    "{{ var }} — insert a value at this spot",
    "{% if data.notes %} ... {% endif %} — show a block conditionally",
    "{%p for item in items %} ... {%p endfor %} — repeat whole paragraphs (one per item)",
    "{%tr for item in items %} ... {%tr endfor %} — repeat a table row (use inside a 1-row table)",
  ]
  for line in syntax_lines:
    doc.add_paragraph(line, style="List Bullet")

  doc.add_heading("Comment-driven placeholders (Word)", level=1)
  doc.add_paragraph(
    "Instead of typing Jinja into the document body, you can keep readable text on the page "
    "and put the placeholder in a Word comment on that text."
  )
  doc.add_paragraph(
    "Steps: type readable text (for example Customer Name), highlight it, insert a comment on "
    "the selection, and type the full placeholder there (for example {{ data.customer_name }}). "
    "On export, the highlighted text is replaced by that placeholder before rendering."
  )
  doc.add_paragraph(
    "Resolved / completed comments: if you mark the comment as done in Word, the placeholder "
    "is skipped and the original highlighted text stays in the export."
  )
  doc.add_paragraph(
    "Loops and conditionals ({%p for ... %}, {% if ... %}) should stay inline in the document "
    "body — do not put those inside comments."
  )

  doc.add_heading("Global placeholders", level=1)
  globals_table = [
    ("{{ account.name }}", "Account name"),
    ("{{ section.label }}", "Section name"),
    ("{{ section.slug }}", "Section slug"),
    ("{{ section.detail }}", "Section detail/subtitle"),
    ("{{ exported_at }}", "Timestamp the PDF was generated"),
    ("{{ exported_by }}", "User who triggered the export"),
  ]
  table = doc.add_table(rows=1, cols=2)
  table.style = "Table Grid"
  hdr = table.rows[0].cells
  hdr[0].text = "Placeholder"
  hdr[1].text = "Description"
  for token, desc in globals_table:
    row = table.add_row().cells
    row[0].text = token
    row[1].text = desc

  doc.add_heading("Single-item shortcuts", level=1)
  doc.add_paragraph(
    "When exporting one item these aliases resolve to that item; "
    "when exporting many they resolve to the first item."
  )
  for token, desc in [
    ("{{ item.name }}", "Item name (object alias)"),
    ("{{ item.created_at }}", "Item created date (object alias)"),
    ("{{ item.data.<field_key> }}", "Item field value (object alias)"),
    ("{{ name }}", "Item name"),
    ("{{ created_at }}", "Item created date"),
    ("{{ data.<field_key> }}", "Any field on the item — see your sections below"),
  ]:
    p = doc.add_paragraph()
    p.add_run(token).bold = True
    p.add_run(f" — {desc}")

  doc.add_heading("Multi-item loop example", level=1)
  example = doc.add_paragraph()
  example.add_run(
    "{%p for item in items %}\n"
    "Item: {{ item.name }}\n"
    "Created: {{ item.created_at }}\n"
    "Notes: {{ item.data.notes }}\n"
    "{%p endfor %}"
  )

  doc.add_heading("Per-section field placeholders", level=1)
  any_section = False
  for section in sections:
    any_section = True
    label = section.get("label") or section.get("slug") or "Section"
    slug = section.get("slug") or ""
    doc.add_heading(f"{label}", level=2)
    if slug:
      sub = doc.add_paragraph()
      sub.add_run(f"slug: {slug}").italic = True

    fields = _safe_field_keys(section.get("schema"))
    if not fields:
      doc.add_paragraph("(No custom fields defined yet.)")
      continue

    section_table = doc.add_table(rows=1, cols=2)
    section_table.style = "Table Grid"
    hdr = section_table.rows[0].cells
    hdr[0].text = "Placeholder"
    hdr[1].text = "Field"
    for key, friendly in fields:
      row = section_table.add_row().cells
      row[0].text = f"{{{{ data.{key} }}}}"
      row[1].text = friendly

  if not any_section:
    doc.add_paragraph("(This account has no sections yet.)")

  doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
  doc.add_heading("Tips", level=1)
  for tip in [
    "Type each {{ ... }} placeholder in one go — Word can split a token across formatting runs and break rendering (comment-driven placeholders avoid this for body text).",
    "For repeating table rows, put the {%tr for item in items %} on the first cell of the row and {%tr endfor %} on the last cell of the same row.",
    "Unknown placeholders (typos) will return a 400 error at export time so you know to fix them.",
  ]:
    doc.add_paragraph(tip, style="List Bullet")

  buf = io.BytesIO()
  doc.save(buf)
  return buf.getvalue()
